import io
import json
import logging
from core.llm_client import LLMClient

logger = logging.getLogger('app')


class AIGrammarService:
    def __init__(self):
        pass

    def check_grammar(self, text, dialect='en-us', use_premium=False):
        """
        Check grammar and score writing quality.
        Returns (result_dict, error_string).
        result_dict contains 'corrections' list and 'writing_scores' dict.
        """
        prompt = f"""You are an expert grammar checker and writing analyst. Analyze the following text for grammar, spelling, punctuation, and style issues.

Dialect: {dialect}

Text to analyze:
\"\"\"
{text}
\"\"\"

Return a JSON object with exactly this structure:
{{
  "corrections": [
    {{
      "original": "the exact text with the error",
      "suggestion": "the corrected text",
      "type": "grammar|spelling|punctuation|style|clarity|wordiness|passive_voice",
      "explanation": "brief explanation of why this is an error and how to fix it",
      "position": {{
        "start": 0,
        "end": 10
      }}
    }}
  ],
  "writing_scores": {{
    "grammar": 85,
    "fluency": 78,
    "clarity": 82,
    "engagement": 70,
    "delivery": 75
  }},
  "tone": "formal|semi-formal|neutral|semi-casual|casual",
  "readability_score": 65.5
}}

Important rules:
- Each score must be an integer from 0-100
- "position" start/end are character indices in the original text
- Only flag genuine errors or meaningful improvements
- Be precise with positions - they must exactly match the original text
- readability_score should be the Flesch-Kincaid reading ease score (0-100)
- Return ONLY valid JSON, no markdown formatting or extra text"""

        try:
            response_text, error = LLMClient.generate(
                system_prompt=None,
                messages=[{"role": "user", "content": prompt}],
                max_tokens=4096,
                use_premium=use_premium
            )

            if error:
                return None, error

            # Strip markdown code fences if present
            if response_text.startswith('```'):
                lines = response_text.split('\n')
                response_text = '\n'.join(lines[1:])
                if response_text.endswith('```'):
                    response_text = response_text[:-3].strip()

            result = json.loads(response_text)

            # Validate structure
            if 'corrections' not in result:
                result['corrections'] = []
            if 'writing_scores' not in result:
                result['writing_scores'] = {
                    'grammar': 50, 'fluency': 50, 'clarity': 50,
                    'engagement': 50, 'delivery': 50
                }
            if 'tone' not in result:
                result['tone'] = 'neutral'
            if 'readability_score' not in result:
                result['readability_score'] = 50.0

            # Clamp scores to 0-100
            for key in ['grammar', 'fluency', 'clarity', 'engagement', 'delivery']:
                score = result['writing_scores'].get(key, 50)
                result['writing_scores'][key] = max(0, min(100, int(score)))

            return result, None

        except json.JSONDecodeError as e:
            logger.error(f"Grammar check JSON parse error: {e}")
            return None, "Failed to parse AI response"
        except Exception as e:
            logger.error(f"Grammar check error: {e}")
            return None, str(e)

    def fix_all(self, text, corrections):
        """
        Apply all corrections to the text at once.
        Processes corrections from end to start to preserve character positions.
        Returns (corrected_text, error_string).
        """
        try:
            if not corrections:
                return text, None

            # Sort by position start descending so we fix from end to beginning
            sorted_corrections = sorted(
                corrections,
                key=lambda c: c.get('position', {}).get('start', 0),
                reverse=True
            )

            result = text
            for correction in sorted_corrections:
                pos = correction.get('position', {})
                start = pos.get('start')
                end = pos.get('end')
                suggestion = correction.get('suggestion', '')

                if start is not None and end is not None:
                    # Verify the original text matches before replacing
                    original_at_pos = result[start:end]
                    expected_original = correction.get('original', '')

                    if original_at_pos == expected_original:
                        result = result[:start] + suggestion + result[end:]
                    else:
                        # Fall back to simple string replace for the first occurrence
                        if expected_original in result:
                            result = result.replace(expected_original, suggestion, 1)

            return result, None

        except Exception as e:
            logger.error(f"Fix all error: {e}")
            return None, str(e)

    def fix_single(self, text, correction):
        """
        Apply a single correction to the text.
        Returns (corrected_text, error_string).
        """
        try:
            pos = correction.get('position', {})
            start = pos.get('start')
            end = pos.get('end')
            suggestion = correction.get('suggestion', '')
            original = correction.get('original', '')

            if start is not None and end is not None:
                original_at_pos = text[start:end]
                if original_at_pos == original:
                    return text[:start] + suggestion + text[end:], None

            # Fallback: replace first occurrence
            if original and original in text:
                return text.replace(original, suggestion, 1), None

            return text, "Could not locate the text to fix"

        except Exception as e:
            logger.error(f"Fix single error: {e}")
            return None, str(e)


class ProofreaderService:
    """
    Document proofreading service that provides comprehensive analysis
    including overall quality score, error categorization, and corrected text.
    """

    ALLOWED_EXTENSIONS = {'docx', 'txt', 'pdf'}
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10 MB

    def __init__(self):
        pass

    def extract_text_from_file(self, uploaded_file):
        """
        Extract text content from an uploaded file (DOCX, TXT, or PDF).
        Returns (text, error).
        """
        filename = uploaded_file.name.lower()

        if filename.endswith('.txt'):
            return self._extract_txt(uploaded_file)
        elif filename.endswith('.docx'):
            return self._extract_docx(uploaded_file)
        elif filename.endswith('.pdf'):
            return self._extract_pdf(uploaded_file)
        else:
            return None, 'Unsupported file type. Please upload a DOCX, TXT, or PDF file.'

    def _extract_txt(self, uploaded_file):
        """Extract text from a .txt file."""
        try:
            content = uploaded_file.read()
            # Try UTF-8 first, fall back to latin-1
            try:
                text = content.decode('utf-8')
            except UnicodeDecodeError:
                text = content.decode('latin-1')
            return text.strip(), None
        except Exception as e:
            logger.error(f"TXT extraction error: {e}")
            return None, 'Failed to read TXT file.'

    def _extract_docx(self, uploaded_file):
        """Extract text from a .docx file."""
        try:
            from docx import Document
            doc = Document(uploaded_file)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            text = '\n\n'.join(paragraphs)
            if not text.strip():
                return None, 'The uploaded DOCX file appears to be empty.'
            return text, None
        except ImportError:
            logger.error("python-docx not installed")
            return None, 'DOCX processing is not available. Please paste your text instead.'
        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            return None, 'Failed to read DOCX file. Please ensure it is a valid document.'

    def _extract_pdf(self, uploaded_file):
        """Extract text from a PDF file."""
        try:
            import pdfplumber
            text_parts = []
            with pdfplumber.open(uploaded_file) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            text = '\n\n'.join(text_parts)
            if not text.strip():
                return None, 'Could not extract text from the PDF. It may be an image-based PDF.'
            return text, None
        except ImportError:
            logger.error("pdfplumber not installed")
            return None, 'PDF processing is not available. Please paste your text instead.'
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            return None, 'Failed to read PDF file. Please ensure it is a valid document.'

    def proofread(self, text, use_premium=False):
        """
        Perform comprehensive proofreading on the given text.
        Returns (result_dict, error_string).

        result_dict contains:
            - overall_score: int (0-100)
            - error_counts: dict with category counts
            - corrections: list of correction dicts
            - corrected_text: the full corrected version
            - summary: brief summary of the document quality
        """
        prompt = f"""You are an expert professional proofreader. Analyze the following document thoroughly for ALL errors and issues across these categories: grammar, spelling, punctuation, style, clarity, and wordiness.

Document to proofread:
\"\"\"
{text}
\"\"\"

Return a JSON object with exactly this structure:
{{
  "overall_score": 78,
  "summary": "A brief 1-2 sentence summary of the document's overall writing quality and main issues found.",
  "error_counts": {{
    "grammar": 3,
    "spelling": 1,
    "punctuation": 2,
    "style": 4,
    "clarity": 2,
    "wordiness": 1
  }},
  "corrections": [
    {{
      "original": "the exact text with the error",
      "suggestion": "the corrected text",
      "type": "grammar|spelling|punctuation|style|clarity|wordiness",
      "explanation": "Brief explanation of why this is an error and the fix"
    }}
  ],
  "corrected_text": "The entire document text with ALL corrections applied. This must be the complete document, not a summary."
}}

Important rules:
- overall_score is 0-100 where 100 is perfect. Deduct points for each error found.
- error_counts must accurately reflect the number of corrections in each category.
- corrections must list EVERY error found, even minor ones.
- Each correction type must be one of: grammar, spelling, punctuation, style, clarity, wordiness
- corrected_text must be the COMPLETE document with all fixes applied, preserving the original structure and paragraphs.
- Be thorough but do not invent errors that do not exist.
- Return ONLY valid JSON, no markdown formatting or extra text."""

        try:
            response_text, error = LLMClient.generate(
                system_prompt=None,
                messages=[{"role": "user", "content": prompt}],
                max_tokens=8192,
                use_premium=use_premium
            )

            if error:
                return None, error

            # Strip markdown code fences if present
            if response_text.startswith('```'):
                lines = response_text.split('\n')
                response_text = '\n'.join(lines[1:])
                if response_text.endswith('```'):
                    response_text = response_text[:-3].strip()

            result = json.loads(response_text)

            # Validate and sanitize structure
            if 'overall_score' not in result:
                result['overall_score'] = 50
            result['overall_score'] = max(0, min(100, int(result['overall_score'])))

            if 'summary' not in result:
                result['summary'] = 'Proofreading analysis complete.'

            if 'error_counts' not in result:
                result['error_counts'] = {}
            for cat in ['grammar', 'spelling', 'punctuation', 'style', 'clarity', 'wordiness']:
                result['error_counts'].setdefault(cat, 0)
                result['error_counts'][cat] = max(0, int(result['error_counts'][cat]))

            if 'corrections' not in result:
                result['corrections'] = []

            if 'corrected_text' not in result or not result['corrected_text'].strip():
                result['corrected_text'] = text  # Fall back to original

            total_errors = sum(result['error_counts'].values())
            result['total_errors'] = total_errors

            return result, None

        except json.JSONDecodeError as e:
            logger.error(f"Proofreader JSON parse error: {e}")
            return None, "Failed to parse AI response. Please try again."
        except Exception as e:
            logger.error(f"Proofreader error: {e}")
            return None, str(e)

    def generate_corrected_docx(self, corrected_text):
        """
        Generate a DOCX file from corrected text.
        Returns (BytesIO, error).
        """
        try:
            from docx import Document
            from docx.shared import Pt

            doc = Document()
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Calibri'
            font.size = Pt(11)

            # Split text into paragraphs and add each
            paragraphs = corrected_text.split('\n')
            for para_text in paragraphs:
                if para_text.strip():
                    doc.add_paragraph(para_text.strip())
                elif paragraphs.index(para_text) > 0:
                    # Preserve blank lines as empty paragraphs
                    doc.add_paragraph('')

            output = io.BytesIO()
            doc.save(output)
            output.seek(0)
            return output, None

        except ImportError:
            logger.error("python-docx not installed for DOCX generation")
            return None, 'DOCX generation is not available.'
        except Exception as e:
            logger.error(f"DOCX generation error: {e}")
            return None, 'Failed to generate DOCX file.'

    def generate_corrected_txt(self, corrected_text):
        """
        Generate a TXT file from corrected text.
        Returns (BytesIO, error).
        """
        try:
            output = io.BytesIO()
            output.write(corrected_text.encode('utf-8'))
            output.seek(0)
            return output, None
        except Exception as e:
            logger.error(f"TXT generation error: {e}")
            return None, 'Failed to generate TXT file.'
